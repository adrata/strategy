#!/usr/bin/env python3
"""
Final Snyk BGI Transformation Report
Pure black/white/gray branding with correct growth projections and direct value messaging
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import os

def create_final_snyk_documents():
    """Create both Word and PDF with corrected branding and messaging"""
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    
    create_final_word_document(desktop_path)
    create_final_pdf_document(desktop_path)

def create_final_word_document(desktop_path):
    """Create Word document with pure black/white/gray branding"""
    filename = os.path.join(desktop_path, "Snyk_BGI_Transformation_Final.docx")
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Pure black/white/gray branding - NO COLOR
    adrata_black = RGBColor(0, 0, 0)
    adrata_gray = RGBColor(102, 102, 102)
    adrata_light_gray = RGBColor(153, 153, 153)
    
    def set_font_style(run, size=11, bold=False, color=adrata_black):
        run.font.name = 'Inter'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
    
    # TITLE PAGE
    title = doc.add_heading("BUYER GROUP TRANSFORMATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font_style(run, 24, True, adrata_black)
    
    subtitle1 = doc.add_heading("Accelerating Snyk's Growth from $300M to $1B+", 1)
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle1.runs:
        set_font_style(run, 18, False, adrata_gray)
    
    subtitle2 = doc.add_heading("Through Systematic Buyer Group Intelligence", 2)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle2.runs:
        set_font_style(run, 16, False, adrata_gray)
    
    doc.add_paragraph()
    
    # Author info
    author_para = doc.add_paragraph("Dan Mirolli, Head of Revenue")
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author_para.runs:
        set_font_style(run, 12, True, adrata_black)
    
    company_para = doc.add_paragraph("Adrata")
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in company_para.runs:
        set_font_style(run, 12, False, adrata_black)
    
    date_para = doc.add_paragraph("July 2025")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        set_font_style(run, 11, False, adrata_gray)
    
    doc.add_page_break()
    
    # EXECUTIVE SUMMARY
    exec_heading = doc.add_heading("Executive Summary: The Growth Acceleration Opportunity", 1)
    for run in exec_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    para1 = doc.add_paragraph(
        "Snyk has built remarkable momentum: $300M in annual recurring revenue, 3,100+ enterprise customers, "
        "and market leadership in the expanding DevSecOps category. The company is positioned for continued "
        "strong growth, with industry analysts projecting 25% annual growth rates for market leaders."
    )
    for run in para1.runs:
        set_font_style(run, 11, False, adrata_black)
    
    para2 = doc.add_paragraph(
        "However, Snyk's greatest opportunity lies in transforming how enterprise sales engages with increasingly "
        "complex buyer ecosystems. With 45 elite sellers targeting 27,000 enterprise prospects, implementing "
        "Adrata's Buyer Group Intelligence (BGI) platform can accelerate growth velocity from strong to exceptional."
    )
    for run in para2.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # THE TRANSFORMATION OPPORTUNITY
    transformation_title = doc.add_paragraph()
    run = transformation_title.add_run("The Transformation Opportunity:")
    set_font_style(run, 11, True, adrata_black)
    
    bullets = [
        "Transform 27,000 prospect accounts into 180,000+ mapped stakeholder relationships",
        "Accelerate average deal size from $44,000 to $85,000 through multi-stakeholder engagement", 
        "Reduce sales cycle length by 40% with precision stakeholder targeting",
        "Achieve 98% win rates vs. competitors through relationship intelligence",
        "Scale revenue growth from 25% annually to 75%+ through systematic buyer group mastery",
        "Generate $700M+ incremental value over 3 years with 2,720:1 ROI"
    ]
    
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # GROWTH TRAJECTORY ANALYSIS
    growth_heading = doc.add_heading("Growth Trajectory Analysis: Good Growth vs. Exceptional Growth", 1)
    for run in growth_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    growth_para = doc.add_paragraph(
        "Snyk will continue growing—the strategic question is velocity and scale. Analysis reveals two "
        "distinct growth trajectories based on buyer engagement sophistication:"
    )
    for run in growth_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # Growth scenarios table
    scenario_table = doc.add_table(rows=5, cols=4)
    scenario_table.style = 'Table Grid'
    scenario_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    scenario_hdr = scenario_table.rows[0].cells
    scenario_hdr[0].text = 'Growth Scenario'
    scenario_hdr[1].text = 'Year 1 Revenue'
    scenario_hdr[2].text = 'Year 2 Revenue' 
    scenario_hdr[3].text = 'Year 3 Revenue'
    
    # CORRECTED scenario data - both growing, BGI growing faster
    scenario_data = [
        ['Strong Growth (Current Path)', '$375M', '$470M', '$590M'],
        ['Accelerated Growth (With BGI)', '$525M', '$790M', '$1.18B'],
        ['Growth Acceleration Benefit', '+$150M', '+$320M', '+$590M'],
        ['Total Additional Value', '+$1.06B over 3 years', '', '']
    ]
    
    for i, row_data in enumerate(scenario_data, 1):
        row_cells = scenario_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_data):
                row_cells[j].text = cell_data
    
    # INVESTMENT ANALYSIS
    doc.add_page_break()
    investment_heading = doc.add_heading("Investment Analysis: Transformative ROI", 1)
    for run in investment_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    investment_title = doc.add_paragraph()
    run = investment_title.add_run("Strategic Investment: $257,500 annually")
    set_font_style(run, 11, True, adrata_black)
    
    investment_items = [
        "BGI Platform: 50 enterprise licenses @ $5,150 each",
        "Complete stakeholder mapping: 180,000+ professional profiles",
        "Real-time intelligence engine with continuous monitoring",
        "Expert consulting and strategic implementation services"
    ]
    
    for item in investment_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # TRANSFORMATIVE VALUE SECTION
    value_heading = doc.add_heading("The Transformative Value of Buyer Group Intelligence", 1)
    for run in value_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    value_para = doc.add_paragraph(
        "Buyer Group Intelligence fundamentally transforms how enterprise sales operates—moving from "
        "individual contact management to complete stakeholder ecosystem mastery."
    )
    for run in value_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # Core transformations
    transformations_title = doc.add_paragraph()
    run = transformations_title.add_run("Core Transformations:")
    set_font_style(run, 11, True, adrata_black)
    
    transformations = [
        "From Single Contacts to Complete Ecosystems: Map all 6-35+ stakeholders per enterprise account",
        "From Reactive to Predictive: AI-powered stakeholder scoring and decision probability modeling",
        "From Generic to Personalized: Tailored messaging and engagement strategies per stakeholder role",
        "From Sales Tools to Intelligence Platform: Replace 5+ traditional sales tools with unified intelligence",
        "From Individual Performance to Systematic Excellence: Repeatable methodology across entire sales team"
    ]
    
    for transformation in transformations:
        p = doc.add_paragraph(transformation, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # RETURN ANALYSIS
    roi_title = doc.add_paragraph()
    run = roi_title.add_run("Return Analysis:")
    set_font_style(run, 11, True, adrata_black)
    
    roi_items = [
        "Year 1 Growth Acceleration: +$150M additional revenue",
        "3-Year Value Creation: +$1.06B incremental growth", 
        "ROI Multiple: 2,720:1 return on strategic investment",
        "Payback Period: 1.4 days from implementation",
        "Market Position: Sustainable competitive advantage through proprietary buyer intelligence"
    ]
    
    for item in roi_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # STRATEGIC RECOMMENDATION
    recommendation_heading = doc.add_heading("Strategic Recommendation", 1)
    for run in recommendation_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    final_title = doc.add_paragraph()
    run = final_title.add_run("Recommended Action:")
    set_font_style(run, 11, True, adrata_black)
    
    final_para = doc.add_paragraph(
        "Proceed with immediate BGI implementation targeting Q3 2025 launch. This timing positions "
        "Snyk to accelerate from strong growth to exceptional growth while market conditions remain "
        "optimal. The $257,500 investment unlocks $1.06B+ in incremental value creation over 3 years "
        "while establishing a sustainable competitive moat through proprietary buyer intelligence."
    )
    for run in final_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # ABOUT ADRATA
    about_heading = doc.add_heading("About Adrata", 1)
    for run in about_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    about_para = doc.add_paragraph(
        "Adrata transforms B2B sales through buyer group intelligence, serving 500+ enterprise "
        "clients worldwide. Our platform has generated over $15B in incremental revenue for "
        "technology companies through systematic stakeholder engagement and predictive buyer intelligence."
    )
    for run in about_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # Save document
    doc.save(filename)
    print(f"✅ Final Word document created: {filename}")

def create_final_pdf_document(desktop_path):
    """Create PDF with pure black/white/gray branding"""
    filename = os.path.join(desktop_path, "Snyk_BGI_Transformation_Final.pdf")
    
    doc = SimpleDocTemplate(filename, pagesize=letter,
                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                          topMargin=1*inch, bottomMargin=1*inch)
    
    styles = getSampleStyleSheet()
    
    # Pure black/white/gray colors - NO COLOR
    adrata_black = colors.Color(0, 0, 0)
    adrata_gray = colors.Color(0.4, 0.4, 0.4)
    adrata_light_gray = colors.Color(0.7, 0.7, 0.7)
    
    # Custom styles
    title_style = ParagraphStyle(
        'AdrataTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=adrata_black,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'AdrataSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=20,
        textColor=adrata_gray,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    heading_style = ParagraphStyle(
        'AdrataHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=adrata_black,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'AdrataBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
        textColor=adrata_black
    )
    
    bullet_style = ParagraphStyle(
        'AdrataBullet',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leftIndent=20,
        fontName='Helvetica',
        textColor=adrata_black
    )
    
    emphasis_style = ParagraphStyle(
        'AdrataEmphasis',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        fontName='Helvetica-Bold',
        textColor=adrata_black
    )
    
    # Content
    content = []
    
    # Title page
    content.append(Spacer(1, 1*inch))
    content.append(Paragraph("BUYER GROUP TRANSFORMATION", title_style))
    content.append(Paragraph("Accelerating Snyk's Growth from $300M to $1B+", subtitle_style))
    content.append(Paragraph("Through Systematic Buyer Group Intelligence", subtitle_style))
    content.append(Spacer(1, 0.5*inch))
    
    # Author info
    content.append(Paragraph("Dan Mirolli, Head of Revenue", ParagraphStyle('Author', parent=styles['Normal'], fontSize=12, alignment=TA_CENTER, fontName='Helvetica-Bold', textColor=adrata_black)))
    content.append(Paragraph("Adrata", ParagraphStyle('Company', parent=styles['Normal'], fontSize=12, alignment=TA_CENTER, fontName='Helvetica', textColor=adrata_black)))
    content.append(Paragraph("July 2025", ParagraphStyle('Date', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER, fontName='Helvetica', textColor=adrata_gray)))
    
    content.append(PageBreak())
    
    # Executive Summary
    content.append(Paragraph("Executive Summary: The Growth Acceleration Opportunity", heading_style))
    content.append(Paragraph(
        "Snyk has built remarkable momentum: $300M in annual recurring revenue, 3,100+ enterprise customers, "
        "and market leadership in the expanding DevSecOps category. The company is positioned for continued "
        "strong growth, with industry analysts projecting 25% annual growth rates for market leaders.",
        body_style
    ))
    
    content.append(Paragraph(
        "However, Snyk's greatest opportunity lies in transforming how enterprise sales engages with increasingly "
        "complex buyer ecosystems. With 45 elite sellers targeting 27,000 enterprise prospects, implementing "
        "Adrata's Buyer Group Intelligence (BGI) platform can accelerate growth velocity from strong to exceptional.",
        body_style
    ))
    
    content.append(Paragraph("<b>The Transformation Opportunity:</b>", emphasis_style))
    content.append(Paragraph("• Transform 27,000 prospect accounts into 180,000+ mapped stakeholder relationships", bullet_style))
    content.append(Paragraph("• Accelerate average deal size from $44,000 to $85,000 through multi-stakeholder engagement", bullet_style))
    content.append(Paragraph("• Reduce sales cycle length by 40% with precision stakeholder targeting", bullet_style))
    content.append(Paragraph("• Achieve 98% win rates vs. competitors through relationship intelligence", bullet_style))
    content.append(Paragraph("• Scale revenue growth from 25% annually to 75%+ through systematic buyer group mastery", bullet_style))
    content.append(Paragraph("• Generate $700M+ incremental value over 3 years with 2,720:1 ROI", bullet_style))
    
    # Growth Trajectory Analysis
    content.append(Paragraph("Growth Trajectory Analysis: Good Growth vs. Exceptional Growth", heading_style))
    content.append(Paragraph(
        "Snyk will continue growing—the strategic question is velocity and scale. Analysis reveals two "
        "distinct growth trajectories based on buyer engagement sophistication:",
        body_style
    ))
    
    # Growth scenarios table - CORRECTED
    scenario_data = [
        ['Growth Scenario', 'Year 1 Revenue', 'Year 2 Revenue', 'Year 3 Revenue'],
        ['Strong Growth (Current Path)', '$375M', '$470M', '$590M'],
        ['Accelerated Growth (With BGI)', '$525M', '$790M', '$1.18B'],
        ['Growth Acceleration Benefit', '+$150M', '+$320M', '+$590M'],
        ['Total Additional Value', '+$1.06B over 3 years', '', '']
    ]
    
    scenario_table = Table(scenario_data, colWidths=[1.8*inch, 1.2*inch, 1.2*inch, 1.2*inch])
    scenario_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), adrata_black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, adrata_black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, adrata_light_gray])
    ]))
    content.append(scenario_table)
    content.append(Spacer(1, 20))
    
    # Investment Analysis
    content.append(PageBreak())
    content.append(Paragraph("Investment Analysis: Transformative ROI", heading_style))
    
    content.append(Paragraph("<b>Strategic Investment: $257,500 annually</b>", emphasis_style))
    content.append(Paragraph("• BGI Platform: 50 enterprise licenses @ $5,150 each", bullet_style))
    content.append(Paragraph("• Complete stakeholder mapping: 180,000+ professional profiles", bullet_style))
    content.append(Paragraph("• Real-time intelligence engine with continuous monitoring", bullet_style))
    content.append(Paragraph("• Expert consulting and strategic implementation services", bullet_style))
    
    # Transformative Value
    content.append(Paragraph("The Transformative Value of Buyer Group Intelligence", heading_style))
    content.append(Paragraph(
        "Buyer Group Intelligence fundamentally transforms how enterprise sales operates—moving from "
        "individual contact management to complete stakeholder ecosystem mastery.",
        body_style
    ))
    
    content.append(Paragraph("<b>Core Transformations:</b>", emphasis_style))
    content.append(Paragraph("• From Single Contacts to Complete Ecosystems: Map all 6-35+ stakeholders per enterprise account", bullet_style))
    content.append(Paragraph("• From Reactive to Predictive: AI-powered stakeholder scoring and decision probability modeling", bullet_style))
    content.append(Paragraph("• From Generic to Personalized: Tailored messaging and engagement strategies per stakeholder role", bullet_style))
    content.append(Paragraph("• From Sales Tools to Intelligence Platform: Replace 5+ traditional sales tools with unified intelligence", bullet_style))
    content.append(Paragraph("• From Individual Performance to Systematic Excellence: Repeatable methodology across entire sales team", bullet_style))
    
    content.append(Paragraph("<b>Return Analysis:</b>", emphasis_style))
    content.append(Paragraph("• Year 1 Growth Acceleration: +$150M additional revenue", bullet_style))
    content.append(Paragraph("• 3-Year Value Creation: +$1.06B incremental growth", bullet_style))
    content.append(Paragraph("• ROI Multiple: 2,720:1 return on strategic investment", bullet_style))
    content.append(Paragraph("• Payback Period: 1.4 days from implementation", bullet_style))
    content.append(Paragraph("• Market Position: Sustainable competitive advantage through proprietary buyer intelligence", bullet_style))
    
    # Strategic Recommendation
    content.append(Paragraph("Strategic Recommendation", heading_style))
    content.append(Paragraph("<b>Recommended Action:</b>", emphasis_style))
    content.append(Paragraph(
        "Proceed with immediate BGI implementation targeting Q3 2025 launch. This timing positions "
        "Snyk to accelerate from strong growth to exceptional growth while market conditions remain "
        "optimal. The $257,500 investment unlocks $1.06B+ in incremental value creation over 3 years "
        "while establishing a sustainable competitive moat through proprietary buyer intelligence.",
        body_style
    ))
    
    # About Adrata
    content.append(Spacer(1, 0.5*inch))
    content.append(Paragraph("About Adrata", ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, textColor=adrata_gray, alignment=TA_CENTER, fontName='Helvetica-Bold')))
    content.append(Paragraph(
        "Adrata transforms B2B sales through buyer group intelligence, serving 500+ enterprise "
        "clients worldwide. Our platform has generated over $15B in incremental revenue for "
        "technology companies through systematic stakeholder engagement and predictive buyer intelligence.",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, textColor=adrata_gray, alignment=TA_CENTER, fontName='Helvetica')
    ))
    
    # Build document
    doc.build(content)
    print(f"✅ Final PDF document created: {filename}")

if __name__ == "__main__":
    create_final_snyk_documents() 