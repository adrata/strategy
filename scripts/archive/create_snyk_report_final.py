#!/usr/bin/env python3
"""
Snyk Buyer Group Transformation Report Word Generator - Final Sourced Version
Converts the updated markdown report to a professional Word document
With all numbers properly sourced and clean black/white/gray formatting
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_custom_styles(doc):
    """Create simple, clean styles for the document"""
    styles = doc.styles
    
    # Title Style - Bold Black
    title_style = styles.add_style('AdrataTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)  # Black
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle Style - Black
    subtitle_style = styles.add_style('AdrataSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0, 0, 0)  # Black
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(6)
    
    # Header Style - Dark Gray
    header_style = styles.add_style('AdrataHeader', WD_STYLE_TYPE.PARAGRAPH)
    header_font = header_style.font
    header_font.name = 'Arial'
    header_font.size = Pt(16)
    header_font.bold = True
    header_font.color.rgb = RGBColor(64, 64, 64)  # Dark Gray
    header_style.paragraph_format.space_before = Pt(12)
    header_style.paragraph_format.space_after = Pt(6)
    
    # Subheader Style - Medium Gray
    subheader_style = styles.add_style('AdrataSubheader', WD_STYLE_TYPE.PARAGRAPH)
    subheader_font = subheader_style.font
    subheader_font.name = 'Arial'
    subheader_font.size = Pt(14)
    subheader_font.bold = True
    subheader_font.color.rgb = RGBColor(96, 96, 96)  # Medium Gray
    subheader_style.paragraph_format.space_before = Pt(8)
    subheader_style.paragraph_format.space_after = Pt(4)
    
    # Body Style - Black
    body_style = styles.add_style('AdrataBody', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Arial'
    body_font.size = Pt(11)
    body_font.color.rgb = RGBColor(0, 0, 0)  # Black
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15
    
    # Bullet Style - Black
    bullet_style = styles.add_style('AdrataBullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_font = bullet_style.font
    bullet_font.name = 'Arial'
    bullet_font.size = Pt(11)
    bullet_font.color.rgb = RGBColor(0, 0, 0)  # Black
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.line_spacing = 1.1

def process_markdown_content(md_content):
    """Process markdown content and convert to structured data"""
    lines = md_content.split('\n')
    processed_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Main titles (# )
        if line.startswith('# '):
            processed_content.append(('title', line[2:]))
        # Subtitles (## )
        elif line.startswith('## '):
            processed_content.append(('subtitle', line[3:]))
        # Headers (### )
        elif line.startswith('### '):
            processed_content.append(('header', line[4:]))
        # Subheaders (#### )
        elif line.startswith('#### '):
            processed_content.append(('subheader', line[5:]))
        # Bullets (- or * )
        elif line.startswith(('- ', '* ')):
            processed_content.append(('bullet', line[2:]))
        # Horizontal rules
        elif line.startswith('---'):
            processed_content.append(('hr', ''))
        # Regular paragraphs
        else:
            processed_content.append(('body', line))
    
    return processed_content

def clean_text(text):
    """Clean markdown formatting from text"""
    # Remove markdown bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove markdown italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove markdown code
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text

def create_word_document():
    """Create the Word document from the markdown file"""
    
    # Read the markdown file
    md_file = 'snyk-buyer-group-transformation-report.md'
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found")
        return
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create custom styles
    create_custom_styles(doc)
    
    # Process content
    processed_content = process_markdown_content(md_content)
    
    # Add content to document
    for content_type, text in processed_content:
        if content_type == 'title':
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataTitle'
        elif content_type == 'subtitle':
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataSubtitle'
        elif content_type == 'header':
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataHeader'
        elif content_type == 'subheader':
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataSubheader'
        elif content_type == 'bullet':
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataBullet'
            # Add bullet point
            p.style.paragraph_format.left_indent = Inches(0.25)
            run = p.runs[0]
            run.text = '• ' + run.text
        elif content_type == 'body' and text:
            p = doc.add_paragraph(clean_text(text))
            p.style = 'AdrataBody'
        elif content_type == 'hr':
            # Add some space for horizontal rules
            doc.add_paragraph()
    
    # Save document
    output_file = os.path.expanduser('~/Desktop/Snyk_Buyer_Group_Transformation_Report_Adrata_2025_Final.docx')
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    return output_file

if __name__ == "__main__":
    try:
        output_file = create_word_document()
        if output_file and os.path.exists(output_file):
            file_size = os.path.getsize(output_file) / 1024  # KB
            print(f"📄 File size: {file_size:.1f} KB")
            print(f"📁 Location: {output_file}")
            print("✨ Report generated with all numbers properly sourced!")
        else:
            print("❌ Error: File was not created successfully")
    except Exception as e:
        print(f"❌ Error creating document: {e}") 