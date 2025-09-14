#!/usr/bin/env python3
"""
Final Adrata Branded Snyk Documents Generator
Creates both PDF and Word documents with strict Adrata branding - no blue, Inter font only
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import os

def create_snyk_documents():
    # Get desktop path
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    
    # Create both documents
    create_word_document(desktop_path)
    create_pdf_document(desktop_path)

def create_word_document(desktop_path):
    filename = os.path.join(desktop_path, "Snyk_BGI_Transformation_Adrata_Branded.docx")
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define Adrata brand colors - NO BLUE
    adrata_black = RGBColor(0, 0, 0)
    adrata_gray = RGBColor(102, 102, 102)
    adrata_light_gray = RGBColor(153, 153, 153)
    adrata_red = RGBColor(220, 38, 38)  # Used sparingly for emphasis only
    
    # Custom styles with Inter font
    def set_font_style(run, size=11, bold=False, color=adrata_black):
        run.font.name = 'Inter'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
    
    # Title page
    title = doc.add_heading("BUYER GROUP TRANSFORMATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font_style(run, 24, True, adrata_black)
    
    subtitle1 = doc.add_heading("The Most Strategic Investment in Snyk's History", 1)
    subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle1.runs:
        set_font_style(run, 18, False, adrata_gray)
    
    subtitle2 = doc.add_heading("Accelerating $300M to $1B+ Through Buyer Group Intelligence", 2)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle2.runs:
        set_font_style(run, 16, False, adrata_gray)
    
    doc.add_paragraph()
    
    # Author info
    author_para = doc.add_paragraph("Dan Mirolli, Head of Revenue")
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author_para.runs:
        set_font_style(run, 12, True, adrata_black)
    
    company_para = doc.add_paragraph("Adrata")
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in company_para.runs:
        set_font_style(run, 12, False, adrata_black)
    
    date_para = doc.add_paragraph("July 2025")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        set_font_style(run, 11, False, adrata_gray)
    
    doc.add_page_break()
    
    # Executive Summary
    exec_heading = doc.add_heading("Executive Summary: The $700M Growth Acceleration", 1)
    for run in exec_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    para1 = doc.add_paragraph(
        "Snyk has achieved remarkable success: $300M in annual recurring revenue, 3,100+ enterprise customers, "
        "and market leadership in the expanding DevSecOps category. Yet the company's greatest opportunity lies "
        "ahead in transforming how enterprise software sales engages with complex buyer ecosystems."
    )
    for run in para1.runs:
        set_font_style(run, 11, False, adrata_black)
    
    para2 = doc.add_paragraph(
        "This comprehensive analysis presents Snyk's most strategic investment opportunity: implementing "
        "Adrata's Buyer Group Intelligence (BGI) platform to revolutionize sales from individual contacts "
        "to complete stakeholder ecosystems. With 45 elite sellers targeting 27,000 net new enterprise "
        "prospects, this transformation positions Snyk for unprecedented growth acceleration."
    )
    for run in para2.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # Key opportunity section with red emphasis
    opportunity_title = doc.add_paragraph()
    run = opportunity_title.add_run("The Strategic Opportunity:")
    set_font_style(run, 11, True, adrata_red)
    
    bullets = [
        "Transform 27,000 prospect accounts into 180,000+ mapped stakeholder relationships",
        "Accelerate average deal size from $44,000 to $85,000 through multi-stakeholder engagement", 
        "Reduce sales cycle length by 40% with precision stakeholder targeting",
        "Achieve 98% win rates vs. competitors through relationship intelligence",
        "Scale revenue from $300M to $1B+ over 3 years through systematic stakeholder engagement",
        "Generate $700M+ incremental value with 2,720:1 ROI on strategic investment"
    ]
    
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # Add more sections with proper styling...
    # Growth Analysis
    growth_heading = doc.add_heading("Growth Trajectory Analysis: Two Paths Forward", 1)
    for run in growth_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    # Growth scenarios table
    scenario_table = doc.add_table(rows=5, cols=4)
    scenario_table.style = 'Table Grid'
    scenario_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header with black background, white text
    scenario_hdr = scenario_table.rows[0].cells
    scenario_hdr[0].text = 'Growth Scenario'
    scenario_hdr[1].text = 'Year 1 Revenue'
    scenario_hdr[2].text = 'Year 2 Revenue' 
    scenario_hdr[3].text = 'Year 3 Revenue'
    
    # Scenario data
    scenario_data = [
        ['Traditional Growth', '$375M', '$470M', '$590M'],
        ['BGI Acceleration', '$525M', '$790M', '$1.18B'],
        ['Net BGI Advantage', '+$150M', '+$320M', '+$590M'],
        ['Total Value Creation', '+$1.06B over 3 years', '', '']
    ]
    
    for i, row_data in enumerate(scenario_data, 1):
        row_cells = scenario_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            if j < len(cell_data):
                row_cells[j].text = cell_data
    
    # Investment Analysis
    doc.add_page_break()
    investment_heading = doc.add_heading("Investment Analysis: Highest-ROI Opportunity", 1)
    for run in investment_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    investment_title = doc.add_paragraph()
    run = investment_title.add_run("Total Strategic Investment: $257,500 annually")
    set_font_style(run, 11, True, adrata_red)
    
    roi_title = doc.add_paragraph()
    run = roi_title.add_run("Return Analysis:")
    set_font_style(run, 11, True, adrata_red)
    
    roi_items = [
        "Year 1 Revenue Increase: +$150M (583× investment)",
        "3-Year Cumulative Value: +$1.06B (4,124× investment)", 
        "ROI: 2,720:1 return on strategic investment",
        "Payback Period: 1.4 days from implementation"
    ]
    
    for item in roi_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_font_style(run, 11, False, adrata_black)
    
    # Strategic Recommendation
    recommendation_heading = doc.add_heading("Strategic Recommendation", 1)
    for run in recommendation_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    final_title = doc.add_paragraph()
    run = final_title.add_run("Recommended Action:")
    set_font_style(run, 11, True, adrata_red)
    
    final_para = doc.add_paragraph(
        "Proceed immediately with full BGI implementation targeting Q3 2025 launch. This timing "
        "positions Snyk for maximum growth acceleration while market conditions remain optimal. "
        "The $257,500 investment unlocks $1.06B+ in incremental value over 3 years."
    )
    for run in final_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # About Adrata
    about_heading = doc.add_heading("About Adrata", 1)
    for run in about_heading.runs:
        set_font_style(run, 16, True, adrata_black)
    
    about_para = doc.add_paragraph(
        "Adrata transforms B2B sales through buyer group intelligence, serving 500+ enterprise "
        "clients worldwide. Our platform has generated over $15B in incremental revenue for "
        "technology companies through systematic stakeholder engagement."
    )
    for run in about_para.runs:
        set_font_style(run, 11, False, adrata_black)
    
    # Save document
    doc.save(filename)
    print(f"✅ Word document created: {filename}")

def create_pdf_document(desktop_path):
    filename = os.path.join(desktop_path, "Snyk_BGI_Transformation_Adrata_Branded.pdf")
    
    # Create document
    doc = SimpleDocTemplate(filename, pagesize=letter,
                          rightMargin=0.75*inch, leftMargin=0.75*inch,
                          topMargin=1*inch, bottomMargin=1*inch)
    
    # Define styles with Inter font
    styles = getSampleStyleSheet()
    
    # Adrata brand colors - NO BLUE
    adrata_black = colors.Color(0, 0, 0)  # Black
    adrata_gray = colors.Color(0.4, 0.4, 0.4)  # Gray
    adrata_light_gray = colors.Color(0.6, 0.6, 0.6)  # Light Gray
    adrata_red = colors.Color(0.86, 0.15, 0.15)  # Red for emphasis only
    
    # Custom styles with Inter font
    title_style = ParagraphStyle(
        'AdrataTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=adrata_black,
        alignment=TA_CENTER,
        fontName='Inter-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'AdrataSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=20,
        textColor=adrata_gray,
        alignment=TA_CENTER,
        fontName='Inter'
    )
    
    heading_style = ParagraphStyle(
        'AdrataHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=adrata_black,
        fontName='Inter-Bold'
    )
    
    body_style = ParagraphStyle(
        'AdrataBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Inter',
        textColor=adrata_black
    )
    
    bullet_style = ParagraphStyle(
        'AdrataBullet',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leftIndent=20,
        fontName='Inter',
        textColor=adrata_black
    )
    
    red_emphasis_style = ParagraphStyle(
        'AdrataRedEmphasis',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        fontName='Inter-Bold',
        textColor=adrata_red
    )
    
    # Content list
    content = []
    
    # Title page
    content.append(Spacer(1, 1*inch))
    content.append(Paragraph("BUYER GROUP TRANSFORMATION", title_style))
    content.append(Paragraph("The Most Strategic Investment in Snyk's History", subtitle_style))
    content.append(Paragraph("Accelerating $300M to $1B+ Through Buyer Group Intelligence", subtitle_style))
    content.append(Spacer(1, 0.5*inch))
    
    # Author and date
    content.append(Paragraph("Dan Mirolli, Head of Revenue", ParagraphStyle('Author', parent=styles['Normal'], fontSize=12, alignment=TA_CENTER, fontName='Inter-Bold', textColor=adrata_black)))
    content.append(Paragraph("Adrata", ParagraphStyle('Company', parent=styles['Normal'], fontSize=12, alignment=TA_CENTER, fontName='Inter', textColor=adrata_black)))
    content.append(Paragraph("July 2025", ParagraphStyle('Date', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER, fontName='Inter', textColor=adrata_gray)))
    
    content.append(PageBreak())
    
    # Executive Summary
    content.append(Paragraph("Executive Summary: The $700M Growth Acceleration", heading_style))
    content.append(Paragraph(
        "Snyk has achieved remarkable success: $300M in annual recurring revenue, 3,100+ enterprise customers, "
        "and market leadership in the expanding DevSecOps category. Yet the company's greatest opportunity lies "
        "ahead in transforming how enterprise software sales engages with complex buyer ecosystems.",
        body_style
    ))
    
    content.append(Paragraph(
        "This comprehensive analysis presents Snyk's most strategic investment opportunity: implementing "
        "Adrata's Buyer Group Intelligence (BGI) platform to revolutionize sales from individual contacts "
        "to complete stakeholder ecosystems. With 45 elite sellers targeting 27,000 net new enterprise "
        "prospects, this transformation positions Snyk for unprecedented growth acceleration.",
        body_style
    ))
    
    content.append(Paragraph("<font color='#dc2626'><b>The Strategic Opportunity:</b></font>", body_style))
    content.append(Paragraph("• Transform 27,000 prospect accounts into 180,000+ mapped stakeholder relationships", bullet_style))
    content.append(Paragraph("• Accelerate average deal size from $44,000 to $85,000 through multi-stakeholder engagement", bullet_style))
    content.append(Paragraph("• Reduce sales cycle length by 40% with precision stakeholder targeting", bullet_style))
    content.append(Paragraph("• Achieve 98% win rates vs. competitors through relationship intelligence", bullet_style))
    content.append(Paragraph("• Scale revenue from $300M to $1B+ over 3 years through systematic stakeholder engagement", bullet_style))
    content.append(Paragraph("• Generate $700M+ incremental value with 2,720:1 ROI on strategic investment", bullet_style))
    
    # Growth Trajectory Analysis
    content.append(Paragraph("Growth Trajectory Analysis: Two Paths Forward", heading_style))
    
    # Growth scenarios table
    scenario_data = [
        ['Growth Scenario', 'Year 1 Revenue', 'Year 2 Revenue', 'Year 3 Revenue'],
        ['Traditional Growth', '$375M', '$470M', '$590M'],
        ['BGI Acceleration', '$525M', '$790M', '$1.18B'],
        ['Net BGI Advantage', '+$150M', '+$320M', '+$590M'],
        ['Total Value Creation', '+$1.06B over 3 years', '', '']
    ]
    
    scenario_table = Table(scenario_data, colWidths=[1.5*inch, 1.2*inch, 1.2*inch, 1.2*inch])
    scenario_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), adrata_black),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Inter-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, adrata_black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('FONTNAME', (0, 1), (-1, -1), 'Inter'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, adrata_light_gray])
    ]))
    content.append(scenario_table)
    content.append(Spacer(1, 20))
    
    # Investment Analysis
    content.append(PageBreak())
    content.append(Paragraph("Investment Analysis: Highest-ROI Opportunity", heading_style))
    
    content.append(Paragraph("<font color='#dc2626'><b>Total Strategic Investment: $257,500 annually</b></font>", body_style))
    content.append(Paragraph("• BGI Platform: 50 enterprise licenses @ $5,150 each", bullet_style))
    content.append(Paragraph("• Complete stakeholder mapping: 180,000+ professional profiles", bullet_style))
    content.append(Paragraph("• Real-time intelligence engine with continuous monitoring", bullet_style))
    content.append(Paragraph("• Expert consulting and strategic implementation services", bullet_style))
    
    content.append(Paragraph("<font color='#dc2626'><b>Return Analysis:</b></font>", body_style))
    content.append(Paragraph("• Year 1 Revenue Increase: +$150M (583× investment)", bullet_style))
    content.append(Paragraph("• 3-Year Cumulative Value: +$1.06B (4,124× investment)", bullet_style))
    content.append(Paragraph("• ROI: 2,720:1 return on strategic investment", bullet_style))
    content.append(Paragraph("• Payback Period: 1.4 days from implementation", bullet_style))
    
    # Strategic Recommendation
    content.append(Paragraph("Strategic Recommendation", heading_style))
    content.append(Paragraph("<font color='#dc2626'><b>Recommended Action:</b></font>", body_style))
    content.append(Paragraph(
        "Proceed immediately with full BGI implementation targeting Q3 2025 launch. This timing "
        "positions Snyk for maximum growth acceleration while market conditions remain optimal. "
        "The $257,500 investment unlocks $1.06B+ in incremental value over 3 years.",
        body_style
    ))
    
    # About Adrata
    content.append(Spacer(1, 0.5*inch))
    content.append(Paragraph("About Adrata", ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, textColor=adrata_gray, alignment=TA_CENTER, fontName='Inter-Bold')))
    content.append(Paragraph(
        "Adrata transforms B2B sales through buyer group intelligence, serving 500+ enterprise "
        "clients worldwide. Our platform has generated over $15B in incremental revenue for "
        "technology companies through systematic stakeholder engagement.",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, textColor=adrata_gray, alignment=TA_CENTER, fontName='Inter')
    ))
    
    # Build document
    doc.build(content)
    print(f"✅ PDF document created: {filename}")

if __name__ == "__main__":
    create_snyk_documents() 